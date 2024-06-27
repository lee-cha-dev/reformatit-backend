import datetime
import io

import uvicorn
import os
import asyncio
import logging
from PIL import Image
import pillow_heif

from fastapi import FastAPI, UploadFile, File, Request, HTTPException, Form
from fastapi.middleware.httpsredirect import HTTPSRedirectMiddleware
from fastapi.responses import FileResponse
from fastapi.responses import JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from slowapi import Limiter, _rate_limit_exceeded_handler
from slowapi.util import get_remote_address

# New imports for document conversion
from docx import Document
from docx2pdf import convert
from odf.opendocument import load
from odf.text import P
from PyPDF2 import PdfReader, PdfWriter

# IMPORTANT TO NOTE THAT THIS THE DOCX CONVERTER NEEDS A REPLACEMENT AS:
# "2024-06-26 22:39:34,298 - main - ERROR - Error during document conversion: docx2pdf is not implemented for linux as
# it requires Microsoft Word to be installed"
# IS THE ERROR PROVIDED. LIBREOFFICE DOES NOT COME INSTALLED ON UBUNTU SO IT IS A POOR CHOICE ON THE UBUNTU SERVER

# INIT APP -
app = FastAPI()

# ... (keep all the existing imports and configurations)
# Consts
# Add these new constants for document conversion
ALLOWED_DOC_FORMATS = ["PDF", "DOCX", "ODT", "TXT", "RTF"]
ALLOWED_DOC_CONTENT_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/vnd.oasis.opendocument.text",
    "text/plain",
    "application/rtf"
}
MAX_DOC_FILE_SIZE = 10 * 1024 * 1024  # 10 MB
TEMP_DOC_DIR = "temp_documents"
if not os.path.exists(TEMP_DOC_DIR):
    os.makedirs(TEMP_DOC_DIR, exist_ok=True)

MAX_FILE_SIZE = 10 * 1024 * 1024
DELETE_IMG_AFTER_SECONDS = 3600  # FOR PROD -> clears images every hour on the server via coroutine
# DELETE_IMG_AFTER_SECONDS = 60  # FOR DEV
TIMEOUT_IMG = 10  # TIMEOUT FOR SERVER REQUESTS
TEMP_IMAGE_DIR = "temp_images"
if not os.path.exists(TEMP_IMAGE_DIR):
    os.makedirs(TEMP_IMAGE_DIR, exist_ok=True)

# Allowed image formats/types
ALLOWED_FORMATS = [
    "BMP",
    "GIF",
    "HEIF",
    "ICO",
    "JPEG",
    "JPG",
    "PCX",
    "PNG",
    "PPM",
    "SGI",
    "WEBP",
    "TIF",
    "TIFF"
]

ALLOWED_CONTENT_TYPES = {
    "image/bmp",
    "image/gif",
    "image/heif",
    "image/vnd.microsoft.icon",
    "image/jpeg",
    "image/pcx",
    "image/png",
    "image/x-portable-pixmap",
    "image/sgi",
    "image/tiff",
    "image/webp"
}

# Register HEIF plugin for heif conversion
pillow_heif.register_heif_opener()

# Configure Logs
log_dir = "logs"
os.makedirs(log_dir, exist_ok=True)
current_time = datetime.datetime.now().strftime("%Y_%m_%d_%H%M%S")
log_file_path = os.path.join(log_dir, current_time + ".log")

file_handler = logging.FileHandler(log_file_path)
stream_handler = logging.StreamHandler()

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    handlers=[file_handler, stream_handler]
)
logger = logging.getLogger(__name__)


# Middleware to log requests and responses
@app.middleware("http")
async def log_requests(request: Request, call_next):
    logger.info(f"Request: {request.method} {request.url}")
    response = await call_next(request)
    logger.info(f"Response status: {response.status_code}")
    return response


# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Allow all origins
    allow_credentials=True,
    allow_methods=["*"],  # Allow all methods
    allow_headers=["*"],  # Allow all headers
)

# Add middleware to redirect HTTP to HTTPS
# app.add_middleware(HTTPSRedirectMiddleware)

# Init Limiter
limiter = Limiter(key_func=get_remote_address)
app.state.limiter = limiter
app.add_exception_handler(429, _rate_limit_exceeded_handler)


# Coroutine to clean up the images on the server periodically - conservation of storage.
async def periodic_cleanup():
    while True:
        now = datetime.datetime.now().timestamp()
        for filename in os.listdir(TEMP_IMAGE_DIR):
            file_path = os.path.join(TEMP_IMAGE_DIR, filename)
            if os.path.isfile(file_path):
                file_creation_time = os.path.getctime(file_path)
                if now - file_creation_time > DELETE_IMG_AFTER_SECONDS:
                    try:
                        os.remove(file_path)
                        logger.info(f"Automatically deleted image file: {file_path}")
                    except Exception as e:
                        logger.error(f"Exception during periodic cleanup: {e}")
        await asyncio.sleep(DELETE_IMG_AFTER_SECONDS)


@app.on_event("startup")
async def startup_event():
    # Ensure temp_images dir exists
    os.makedirs(TEMP_IMAGE_DIR, exist_ok=True)
    # THIS CANNOT BE SET TO AWAIT - THE SERVER WILL HANG
    asyncio.create_task(periodic_cleanup())


@app.post("/test/")
@app.get("/test/")
@limiter.limit("10/minute")
async def test(request: Request):
    print(f"Testing {request.url}")
    response_data = {"message": "This is a test response", "url": str(request.url)}
    return JSONResponse(content=response_data)


@app.post("/convert/")
# @limiter.limit("5/minute")
async def convert_image(request: Request, file: UploadFile = File(...), convert_to: str = Form(...)):
    try:
        ######################################################
        # Cybersecurity policies start here
        ######################################################
        # Check if file size is within limit
        contents = await file.read()
        print(f"Converted to: {convert_to}")
        if len(contents) > MAX_FILE_SIZE:
            logger.warning(f"File too large: {len(contents)}")
            raise HTTPException(status_code=413, detail="File too large")

        # Check for the correct format
        if convert_to.upper() not in ALLOWED_FORMATS:
            logger.info("Invalid conversion format")
            raise HTTPException(status_code=400, detail="Invalid conversion format")

        # REMOVED CONTENT TYPE CHECK DUE TO ISSUE WITH CONTENT TYPE NOT BEING PASSED OVER.

        ######################################################
        # Cybersecurity policies end here
        ######################################################
        # Reset the file pointer after reading contents
        file.file.seek(0)

        # Wrap image conversion logic in async method to enable timeout
        async def get_output_path(convert_img_to: str):
            try:
                if convert_img_to.upper() == "HEIF":
                    # Separate handling for HEIF conversion
                    input_path = os.path.join(TEMP_IMAGE_DIR, file.filename)
                    with open(input_path, "wb") as f:
                        f.write(contents)

                    img = Image.open(input_path)
                    if img.mode != "RGB":
                        img = img.convert("RGB")

                    heif_file = pillow_heif.from_pillow(img)

                    output_dir = "temp_images/"
                    os.makedirs(output_dir, exist_ok=True)

                    timestamp = datetime.datetime.now().strftime("%S%M%H%d%m%Y")
                    base_name = os.path.splitext(file.filename)[0]
                    new_filename = f"{base_name}_{timestamp}.heif"
                    out_path = os.path.join(output_dir, new_filename)

                    heif_file.save(out_path, quality=95)

                    logger.info("Image converted successfully to HEIF")

                    # General handling for all other formats
                else:
                    # Correct format strings for certain formats
                    convert_img_to_corrected = {
                        "JPG": "JPEG",
                        "TIF": "TIFF"
                    }.get(convert_img_to.upper(), convert_img_to.upper())

                    # Open and convert the image
                    img = Image.open(file.file)
                    old_format = img.format
                    img = img.convert("RGB")

                    # Ensure temp_images dir exists
                    output_dir = "temp_images/"
                    os.makedirs(output_dir, exist_ok=True)

                    # Generate a timestamp
                    timestamp = datetime.datetime.now().strftime("%S%M%H%d%m%Y")
                    base_name = os.path.splitext(file.filename)[0]
                    new_filename = f"{base_name}_{timestamp}.{convert_img_to_corrected.lower()}"

                    # Save the image to the temporary image directory
                    out_path = os.path.join(output_dir, new_filename)

                    img.save(out_path, format=convert_img_to_corrected)

                    logger.info(
                        f"Image converted successfully. Image converted from {old_format} to {convert_img_to_corrected}")
                    print(f"Output path: {out_path}")

                # Debug statement to verify path
                if not os.path.exists(out_path):
                    raise Exception(f"Failed to save the image. Path does not exist: {out_path}")

                # Check if the file exists and has non-zero size
                if os.path.exists(out_path) and os.path.getsize(out_path) > 0:
                    logger.info(f"File successfully saved at: {out_path}")
                else:
                    logger.error(f"File was not saved correctly at: {out_path}")
                    raise Exception(f"File was not saved correctly at: {out_path}")

                return out_path
            except Exception as e:
                logger.error(f"Error during conversion: {e}\t {file.content_type} to {convert_to}")
                raise HTTPException(status_code=500, detail=str(e))

        # Get the output path with an expected timeout for the server request.
        output_path = await asyncio.wait_for(get_output_path(convert_to), timeout=TIMEOUT_IMG)

        return FileResponse(output_path, media_type=f"image/{convert_to.lower()}")
    except asyncio.TimeoutError:
        logger.error(f"Timeout while converting image")
        raise HTTPException(status_code=504, detail="Timeout while converting image")


@app.post("/convert-document/")
@limiter.limit("5/minute")
async def convert_document(request: Request, file: UploadFile = File(...), convert_to: str = Form(...)):
    try:
        # Check if file size is within limit
        contents = await file.read()
        if len(contents) > MAX_DOC_FILE_SIZE:
            logger.warning(f"Document too large: {len(contents)}")
            raise HTTPException(status_code=413, detail="Document too large")

        # Check for the correct format
        if convert_to.upper() not in ALLOWED_DOC_FORMATS:
            logger.info("Invalid document conversion format")
            raise HTTPException(status_code=400, detail="Invalid document conversion format")

        # Reset the file pointer after reading contents
        file.file.seek(0)

        # Wrap document conversion logic in async method to enable timeout
        async def convert_doc():
            try:
                input_path = os.path.join(TEMP_DOC_DIR, file.filename)
                with open(input_path, "wb") as f:
                    f.write(contents)

                timestamp = datetime.datetime.now().strftime("%S%M%H%d%m%Y")
                base_name = os.path.splitext(file.filename)[0]
                new_filename = f"{base_name}_{timestamp}.{convert_to.lower()}"
                output_path_doc = os.path.join(TEMP_DOC_DIR, new_filename)

                # Determine input format
                input_format = os.path.splitext(file.filename)[1][1:].upper()

                # Perform conversion based on input and output formats
                if input_format == "DOCX" and convert_to.upper() == "PDF":
                    convert(input_path, output_path_doc)
                elif input_format == "PDF" and convert_to.upper() == "DOCX":
                    pdf = PdfReader(input_path)
                    doc = Document()
                    for page in pdf.pages:
                        doc.add_paragraph(page.extract_text())
                    doc.save(output_path_doc)
                elif input_format == "ODT" and convert_to.upper() in ["DOCX", "PDF", "TXT"]:
                    doc = load(input_path)
                    if convert_to.upper() == "DOCX":
                        docx = Document()
                        for element in doc.getElementsByType(P):
                            docx.add_paragraph(str(element))
                        docx.save(output_path_doc)
                    elif convert_to.upper() == "PDF":
                        # First convert to DOCX, then to PDF
                        temp_docx = os.path.join(TEMP_DOC_DIR, f"temp_{timestamp}.docx")
                        docx = Document()
                        for element in doc.getElementsByType(P):
                            docx.add_paragraph(str(element))
                        docx.save(temp_docx)
                        convert(temp_docx, output_path_doc)
                        os.remove(temp_docx)
                    elif convert_to.upper() == "TXT":
                        with open(output_path_doc, 'w', encoding='utf-8') as f:
                            for element in doc.getElementsByType(P):
                                f.write(str(element) + '\n')
                elif convert_to.upper() == "TXT":
                    # For any format to TXT, we'll use a simple text extraction
                    if input_format == "PDF":
                        pdf = PdfReader(input_path)
                        with open(output_path_doc, 'w', encoding='utf-8') as f:
                            for page in pdf.pages:
                                f.write(page.extract_text())
                    elif input_format in ["DOCX", "RTF"]:
                        doc = Document(input_path)
                        with open(output_path_doc, 'w', encoding='utf-8') as f:
                            for para in doc.paragraphs:
                                f.write(para.text + '\n')
                else:
                    raise ValueError(f"Conversion from {input_format} to {convert_to} is not supported.")

                logger.info(f"Document converted successfully to {convert_to}")

                if not os.path.exists(output_path_doc):
                    raise Exception(f"Failed to save the document. Path does not exist: {output_path_doc}")

                if os.path.exists(output_path_doc) and os.path.getsize(output_path_doc) > 0:
                    logger.info(f"Document successfully saved at: {output_path_doc}")
                else:
                    logger.error(f"Document was not saved correctly at: {output_path_doc}")
                    raise Exception(f"Document was not saved correctly at: {output_path_doc}")

                return output_path_doc
            except Exception as e:
                logger.error(f"Error during document conversion: {e}")
                raise HTTPException(status_code=500, detail=str(e))

        # Get the output path with an expected timeout for the server request.
        output_path = await asyncio.wait_for(convert_doc(), timeout=TIMEOUT_IMG)

        return FileResponse(output_path, media_type=f"application/{convert_to.lower()}")
    except asyncio.TimeoutError:
        logger.error(f"Timeout while converting document")
        raise HTTPException(status_code=504, detail="Timeout while converting document")


if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8080)
